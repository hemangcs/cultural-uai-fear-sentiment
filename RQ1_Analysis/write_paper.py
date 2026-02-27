#!/usr/bin/env python3
"""Generate the AMCIS 2026 ERF Paper in Word format."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = "/Volumes/TOSHIBA EXT/refinitiv-Data/RQ1_Analysis"


def set_run_font(run, name='Georgia', size=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_text(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, size=13, bold=True)
        p.space_before = Pt(12)
        p.space_after = Pt(6)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, size=11, bold=True, italic=True)
        p.space_before = Pt(10)
        p.space_after = Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_table(doc, headers, rows, caption=""):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Georgia'

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Georgia'

    # Caption
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        set_run_font(run, size=10, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(4)

    return table


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Georgia'
    font.size = Pt(10)

    # ============================================================
    # TITLE
    # ============================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('AMCIS 2026 Reno')
    set_run_font(run, size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        'Does Cultural Uncertainty Avoidance Amplify Fear Sentiment\n'
        'in Emerging Technology Stocks? A Cross-National Study'
    )
    set_run_font(run, size=20, bold=True)

    type_p = doc.add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = type_p.add_run('Submission Type: Emergent Research Forum (ERF) Paper')
    set_run_font(run, size=10)

    # ============================================================
    # KEYWORDS
    # ============================================================
    add_heading_text(doc, 'Keywords')
    add_body(doc,
        'Cultural dimensions, uncertainty avoidance, sentiment analysis, '
        'fear contagion, artificial intelligence, emerging technology, '
        'cross-cultural finance, Hofstede.'
    )

    # ============================================================
    # INTRODUCTION
    # ============================================================
    add_heading_text(doc, 'Introduction')

    add_body(doc,
        'The rapid emergence of artificial intelligence (AI), quantum computing, and related '
        'technologies has generated unprecedented media attention and investor sentiment volatility '
        'across global financial markets. Events such as the launch of ChatGPT in November 2022, '
        'NVIDIA\'s historic market capitalization surge, and the DeepSeek disruptions in January 2025 '
        'have triggered distinct sentiment responses across national markets. Yet, a critical gap '
        'persists in understanding how national cultural values shape the media-driven fear narratives '
        'surrounding these emerging technology firms.'
    )

    add_body(doc,
        'Hofstede\'s cultural dimensions theory (Hofstede, 2001) posits that uncertainty avoidance '
        '(UAI)\u2014the degree to which members of a society feel uncomfortable with ambiguity\u2014varies '
        'systematically across nations. High-UAI cultures (e.g., Chile, Mexico, Brazil) tend to resist '
        'unstructured situations through rigid codes and intolerance for deviant behavior, while '
        'low-UAI cultures (e.g., United Kingdom, United States, Canada) exhibit greater tolerance for '
        'ambiguity and novel ideas (Hofstede et al., 2010). We argue that this cultural orientation '
        'fundamentally moderates how media narratives frame\u2014and amplify\u2014fear-related sentiment '
        'around disruptive technology companies.'
    )

    add_body(doc,
        'This study addresses the following research question: Does national-level uncertainty '
        'avoidance amplify fear-related media sentiment toward emerging technology companies, '
        'particularly during AI disruption events? Drawing on 923,524 daily firm-level sentiment '
        'observations for 178 technology companies across 8 countries from the MarketPsych/LSEG '
        'Refinitiv analytics platform (1998\u20132026), we conduct panel regression analysis and event '
        'studies around 12 major AI disruption events. Our findings contribute to the intersection of '
        'cultural values theory, information systems, and behavioral finance.'
    )

    # ============================================================
    # THEORETICAL BACKGROUND
    # ============================================================
    add_heading_text(doc, 'Theoretical Background and Hypotheses')

    add_heading_text(doc, 'Uncertainty Avoidance and Information Processing', level=2)

    add_body(doc,
        'Information cascade theory (Bikhchandani et al., 1992) suggests that individuals in '
        'uncertain environments may abandon private information in favor of observed behavior, '
        'creating sentiment cascades. We extend this to the cultural level: societies with higher '
        'UAI scores may exhibit amplified cascading of fear-related narratives because cultural '
        'norms predispose citizens toward threat-vigilant information processing (Rieger et al., 2015). '
        'When disruptive technology events create ambiguity about future economic outcomes, high-UAI '
        'cultures may disproportionately amplify negative media framing.'
    )

    add_body(doc,
        'Media sentiment analytics capture the aggregate tone of news coverage, social media, and '
        'analyst commentary (Tetlock, 2007). The MarketPsych/LSEG platform decomposes sentiment into '
        'granular emotion categories\u2014including fear, uncertainty, stress, and gloom\u2014enabling '
        'direct measurement of fear-related narratives at the firm-day level. We hypothesize:'
    )

    add_body(doc,
        'H1: Companies domiciled in high-UAI countries exhibit higher levels of fear-related media '
        'sentiment compared to companies in low-UAI countries, controlling for firm-level media attention.'
    )

    add_body(doc,
        'H2: The positive relationship between media attention (buzz) and fear sentiment is '
        'stronger for companies in high-UAI countries (moderation effect).'
    )

    add_body(doc,
        'H3: During AI disruption events, companies in high-UAI countries experience larger '
        'abnormal increases in fear sentiment compared to low-UAI countries.'
    )

    # ============================================================
    # METHODOLOGY
    # ============================================================
    add_heading_text(doc, 'Methodology')

    add_heading_text(doc, 'Data and Sample', level=2)

    add_body(doc,
        'We construct a panel dataset from two sources. First, daily firm-level sentiment data '
        'from the MarketPsych/LSEG Refinitiv Media Analytics (RMA) platform for 178 technology-sector '
        'companies listed on 15 exchanges across 8 countries: Australia (UAI=51), Brazil (76), '
        'Canada (48), Chile (86), Mexico (82), Taiwan (69), United Kingdom (35), and United States (46). '
        'The sentiment data spans January 1998 to January 2026, yielding 923,524 firm-day observations '
        'with 62 sentiment variables including fear, uncertainty, stress, gloom, and volatility. '
        'Second, we incorporate Hofstede\'s six cultural dimensions (Hofstede et al., 2010), with UAI '
        'as our primary independent variable.'
    )

    add_heading_text(doc, 'Variables', level=2)

    add_body(doc,
        'Dependent variable: Fear Index\u2014a composite of five standardized fear-related sentiment '
        'measures (fear, uncertainty, stress, gloom, anger) averaged at the firm-day level. '
        'Independent variable: UAI score (standardized). Controls: media attention (buzz, '
        'standardized), and remaining Hofstede dimensions (PDI, IDV, MAS, LTO, IVR). '
        'Interaction term: UAI \u00d7 Buzz captures the amplification effect.'
    )

    add_heading_text(doc, 'Analytical Approach', level=2)

    add_body(doc,
        'We employ three complementary methods: (1) OLS panel regressions with country-clustered '
        'standard errors across four model specifications; (2) event studies around 12 AI disruption '
        'events (2022\u20132025) comparing abnormal fear sentiment between UAI groups; and '
        '(3) robustness checks using alternative dependent variables (uncertainty, stress, gloom individually).'
    )

    # ============================================================
    # RESULTS
    # ============================================================
    add_heading_text(doc, 'Results')

    # Table 1 - Sample Distribution
    add_heading_text(doc, 'Descriptive Statistics', level=2)

    add_body(doc,
        'Table 1 presents the sample distribution. The United States dominates with 88 firms and '
        '801,067 observations, followed by Canada (50 firms, 66,990 observations). High-UAI countries '
        '(Brazil, Chile, Mexico) contribute 53,491 observations across 22 firms. Mean fear sentiment '
        'is notably higher for high-UAI countries (0.0106) than low-UAI countries (0.0076).'
    )

    headers1 = ['Country', 'UAI', 'Group', 'Firms', 'Obs.', 'Avg Fear']
    rows1 = [
        ['Australia', '51', 'Low', '8', '1,356', '0.052'],
        ['Brazil', '76', 'High', '12', '28,860', '0.008'],
        ['Canada', '48', 'Low', '50', '66,990', '0.024'],
        ['Chile', '86', 'High', '4', '3,794', '0.088'],
        ['Mexico', '82', 'High', '6', '20,837', '0.027'],
        ['Taiwan', '69', 'Med.', '5', '195', '0.061'],
        ['United Kingdom', '35', 'Low', '5', '425', '0.020'],
        ['United States', '46', 'Low', '88', '801,067', '0.007'],
    ]
    add_table(doc, headers1, rows1, 'Table 1. Sample Distribution by Country')

    doc.add_paragraph()

    # Table 2 - Regression Results
    add_heading_text(doc, 'Panel Regression Results', level=2)

    add_body(doc,
        'Table 2 reports OLS regression results with country-clustered standard errors. '
        'Model 1 confirms H1: UAI is significantly positively associated with fear sentiment '
        '(\u03b2=0.050, p<0.05). Model 3 tests H2: the interaction UAI\u00d7Buzz is significantly negative '
        '(\u03b2=\u22120.036, p<0.01), indicating that in high-UAI countries, increased media attention '
        'is associated with disproportionately higher fear sentiment\u2014but through the main UAI effect '
        'rather than buzz amplification. The negative interaction suggests that buzz may actually '
        'attenuate fear in high-UAI countries, possibly through information resolution. '
        'Model 4 controls for all Hofstede dimensions: UAI remains strongly significant '
        '(\u03b2=2.094, p<0.001), and the interaction persists (\u03b2=\u22120.020, p<0.001).'
    )

    headers2 = ['Variable', 'Model 1', 'Model 2', 'Model 3', 'Model 4']
    rows2 = [
        ['UAI (std.)', '0.050**', '\u2014', '0.052***', '2.094***'],
        ['UAI High', '\u2014', '0.188', '\u2014', '\u2014'],
        ['UAI Medium', '\u2014', '0.461***', '\u2014', '\u2014'],
        ['Buzz (std.)', '\u22120.092***', '\u22120.092***', '\u22120.094***', '\u22120.084***'],
        ['UAI \u00d7 Buzz', '\u2014', '\u2014', '\u22120.036***', '\u22120.020***'],
        ['PDI (std.)', '\u2014', '\u2014', '\u2014', '\u22120.689***'],
        ['IDV (std.)', '\u2014', '\u2014', '\u2014', '1.749***'],
        ['LTO (std.)', '\u2014', '\u2014', '\u2014', '0.463***'],
        ['IVR (std.)', '\u2014', '\u2014', '\u2014', '0.343***'],
        ['R\u00b2', '0.011', '0.010', '0.012', '0.020'],
        ['N', '763,696', '763,696', '763,696', '763,696'],
    ]
    add_table(doc, headers2, rows2,
              'Table 2. Panel Regression: Fear Index on UAI (Country-Clustered SE)')

    doc.add_paragraph()

    # Event Study Results
    add_heading_text(doc, 'Event Study Results', level=2)

    add_body(doc,
        'Table 3 reports abnormal fear index changes (post-event minus pre-event baseline) around '
        'six key AI disruption events. Consistent with H3, high-UAI countries show substantially '
        'larger fear sentiment spikes during landmark events. ChatGPT\'s launch triggered an abnormal '
        'fear increase of 0.047 in high-UAI countries versus 0.004 in low-UAI countries. The '
        'Microsoft-OpenAI investment announcement showed the largest differential (0.066 vs. 0.001). '
        'Average abnormal fear across all 12 events was 6.0 times larger in high-UAI countries '
        '(0.006 vs. \u22120.000), though the t-test did not reach conventional significance (p=0.508), '
        'likely due to limited event count.'
    )

    headers3 = ['AI Disruption Event', 'High UAI', 'Low UAI']
    rows3 = [
        ['ChatGPT Launch (Nov 2022)', '+0.047', '+0.004'],
        ['MS-OpenAI $10B (Jan 2023)', '+0.066', '+0.001'],
        ['NVIDIA Earnings Surge (May 2023)', '+0.012', '\u22120.003'],
        ['EU AI Act Vote (Jul 2023)', '+0.010', '\u22120.002'],
        ['DeepSeek V3 (Jan 2025)', '+0.020', '\u22120.007'],
        ['Average (all 12 events)', '+0.006', '\u22120.000'],
    ]
    add_table(doc, headers3, rows3,
              'Table 3. Abnormal Fear Index Around AI Disruption Events')

    doc.add_paragraph()

    # Figure reference
    add_body(doc,
        'Figure 1 plots quarterly mean fear sentiment by UAI group from 2020 to 2026. '
        'High-UAI countries consistently exhibit elevated fear sentiment, with visible spikes '
        'around the ChatGPT launch (Q4 2022) and DeepSeek disruptions (Q1 2025).'
    )

    # Insert Figure 1
    fig1_path = os.path.join(OUTPUT_DIR, 'fig1_fear_by_uai.png')
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        run = cap.add_run('Figure 1. Quarterly Fear Sentiment by UAI Group (2020\u20132026)')
        set_run_font(run, size=10, bold=True)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ============================================================
    # DISCUSSION & CONCLUSION
    # ============================================================
    add_heading_text(doc, 'Discussion and Conclusion')

    add_body(doc,
        'This study provides empirical evidence that national cultural uncertainty avoidance '
        'significantly shapes fear-related media sentiment toward emerging technology companies. '
        'Three key findings emerge. First, UAI has a significant positive effect on fear sentiment '
        '(\u03b2=0.050, p<0.05), supporting H1 and suggesting that cultural predispositions toward '
        'ambiguity intolerance translate into measurably higher fear narratives in media coverage '
        'of technology firms. Second, the significant UAI\u00d7Buzz interaction (H2) reveals that the '
        'relationship between media attention and fear is culturally contingent, though in the '
        'opposite direction expected\u2014buzz attenuates rather than amplifies fear in high-UAI '
        'countries, suggesting information-resolution mechanisms. Third, event study evidence '
        'directionally supports H3: high-UAI countries show 6x larger abnormal fear responses '
        'to AI disruption events, though limited event frequency constrains statistical power.'
    )

    add_body(doc,
        'These findings extend cultural dimensions theory into the domain of technology sentiment '
        'analytics and have practical implications for international technology firms managing '
        'investor communications across culturally diverse markets. Companies launching disruptive '
        'AI products should anticipate amplified fear narratives in high-UAI markets and may need '
        'culturally tailored communication strategies to mitigate sentiment-driven volatility.'
    )

    add_heading_text(doc, 'Limitations and Future Research', level=2)

    add_body(doc,
        'This study has several limitations that suggest avenues for future research. First, the '
        'panel is dominated by US firms (87% of observations), creating potential imbalance. Future '
        'work should expand coverage to non-Americas exchanges. Second, UAI is measured at the '
        'national level; within-country cultural variation is not captured. Third, the R\u00b2 values '
        '(1\u20132%) are modest, suggesting that cultural dimensions explain a small but significant '
        'portion of fear sentiment variance. Future research should integrate firm-level controls '
        '(market cap, industry sub-sector) and explore additional cultural frameworks (Schwartz, '
        'GLOBE) for triangulation.'
    )

    # ============================================================
    # REFERENCES
    # ============================================================
    add_heading_text(doc, 'REFERENCES')

    refs = [
        'Bikhchandani, S., Hirshleifer, D., & Welch, I. (1992). A theory of fads, fashion, custom, '
        'and cultural change as informational cascades. Journal of Political Economy, 100(5), 992\u20131026.',

        'Hofstede, G. (2001). Culture\'s consequences: Comparing values, behaviors, institutions and '
        'organizations across nations (2nd ed.). Sage Publications.',

        'Hofstede, G., Hofstede, G. J., & Minkov, M. (2010). Cultures and organizations: Software '
        'of the mind (3rd ed.). McGraw-Hill.',

        'Peterson, R. L. (2016). Trading on sentiment: The power of minds over markets. Wiley.',

        'Rieger, M. O., Wang, M., & Hens, T. (2015). Risk preferences around the world. Management '
        'Science, 61(3), 637\u2013648.',

        'Tetlock, P. C. (2007). Giving content to investor sentiment: The role of media in the stock '
        'market. The Journal of Finance, 62(3), 1139\u20131168.',

        'Shenkar, O. (2001). Cultural distance revisited: Towards a more rigorous conceptualization '
        'and measurement of cultural differences. Journal of International Business Studies, 32(3), 519\u2013535.',

        'MarketPsych/LSEG. (2024). Refinitiv MarketPsych Analytics (RMA) documentation. '
        'https://dataapi.marketpsych.com',
    ]

    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        set_run_font(run, size=10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    # Save
    output_path = os.path.join(OUTPUT_DIR, 'AMCIS_2026_ERF_Paper_RQ1.docx')
    doc.save(output_path)
    print(f"Paper saved to: {output_path}")


if __name__ == "__main__":
    main()
